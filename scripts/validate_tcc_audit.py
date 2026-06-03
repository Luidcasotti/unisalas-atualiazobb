from pathlib import Path
import zipfile

import fitz
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Auditoria_TCC_vs_Codigo_Capitulos_Corrigidos.docx"
PDF = ROOT / "docs" / "Auditoria_TCC_vs_Codigo_Capitulos_Corrigidos.pdf"
PNG_DIR = ROOT / "docs" / "audit_pdf_preview"


def main():
    doc = Document(DOCX)
    with zipfile.ZipFile(DOCX) as zf:
        names = set(zf.namelist())
        for required in ["[Content_Types].xml", "word/document.xml", "word/styles.xml"]:
            if required not in names:
                raise RuntimeError(f"DOCX sem parte obrigatória: {required}")

    pdf = fitz.open(PDF)
    PNG_DIR.mkdir(exist_ok=True)
    for page_index in range(len(pdf)):
        pix = pdf.load_page(page_index).get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
        pix.save(PNG_DIR / f"page-{page_index + 1}.png")

    print(f"DOCX bytes: {DOCX.stat().st_size}")
    print(f"DOCX paragraphs: {len(doc.paragraphs)}")
    print(f"DOCX tables: {len(doc.tables)}")
    print(f"PDF bytes: {PDF.stat().st_size}")
    print(f"PDF pages: {len(pdf)}")
    print(f"Preview: {PNG_DIR}")


if __name__ == "__main__":
    main()
