from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "Auditoria_TCC_vs_Codigo_Capitulos_Corrigidos.docx"
PDF_PATH = OUT_DIR / "Auditoria_TCC_vs_Codigo_Capitulos_Corrigidos.pdf"
TODAY = date.today().strftime("%d/%m/%Y")


DIVERGENCES = [
    ["D01", "Resumo", "Nomeia a plataforma como Uniorte Salas/Uni Salas, enquanto o sistema, layout e documentação técnica usam UniSalas.", "Padronizar o nome como UniSalas em todo o texto."],
    ["D02", "Resumo/Abstract", "Afirma consulta de capacidade, recursos audiovisuais e horários ocupados. O código cadastra observação da sala e tem coluna capacidade na migration, mas o model/view não usam capacidade nem recursos audiovisuais como campos próprios.", "Descrever que o sistema cadastra nome, bloco, cor, observação e manutenção, sem afirmar recursos audiovisuais estruturados."],
    ["D03", "Resumo", "Afirma automação integral e reserva autônoma. No código, muitas reservas ficam pendentes ou em análise e dependem de aprovação administrativa.", "Explicar que o processo é semiautomatizado: solicitação pelo professor, validação de conflito e decisão administrativa."],
    ["D04", "Resumo", "Afirma disponibilidade em tempo real de salas livres. O sistema verifica disponibilidade por sala, data e período selecionados; não há calendário geral em tempo real.", "Substituir por verificação de disponibilidade no momento da solicitação."],
    ["D05", "Introdução", "Afirma consulta a capacidade e recursos. A implementação usa observação textual, não cadastro estruturado de recursos.", "Citar observações da sala, bloco e status de manutenção."],
    ["D06", "Levantamento de requisitos", "Lista edição de reservas como requisito funcional. O código possui tela professor/editar-reserva.blade.php estática, sem rota funcional de edição.", "Remover edição de reservas ou marcar como funcionalidade futura; manter desistência/cancelamento."],
    ["D07", "Levantamento de requisitos", "Não lista funcionalidades já implementadas: avisos, mensagens diretas, manutenção, recorrência, aprovação em lote e notificações.", "Incluir esses requisitos funcionais no capítulo."],
    ["D08", "Estrutura do sistema", "Descreve o administrador apenas como responsável por salas e agendamentos. O código implementa usuários, blocos, salas, manutenção, reservas, avisos, mensagens e histórico.", "Ampliar a descrição do perfil administrador."],
    ["D09", "Estrutura do sistema", "Descreve o professor apenas como usuário que visualiza salas e reserva. O código implementa painel, minhas reservas, histórico, recorrência, desistência e mensagens.", "Ampliar a descrição do perfil professor."],
    ["D10", "Estrutura do sistema", "Usa tempo futuro, como 'será projetada', apesar de o sistema já estar implementado.", "Reescrever no passado/presente, indicando sistema desenvolvido."],
    ["D11", "Desenvolvimento e implementação", "O texto é genérico e não cita Laravel, PHP, Blade, Bootstrap, MySQL/MariaDB, Vite, PHPUnit nem a arquitetura MVC real.", "Inserir seção técnica com tecnologias e organização MVC."],
    ["D12", "Problematização / Tabela de soluções", "Afirma confirmação automática da reserva. No código, reserva pode ficar pendente, em análise, aprovada automaticamente em situações específicas ou cancelada por conflito/manutenção.", "Explicar status reais: pendente, em_analise, aprovada, rejeitada e cancelada."],
    ["D13", "Problematização / Tabela de soluções", "Afirma módulo administrativo com relatórios. O sistema tem dashboard, contadores e histórico filtrável, mas não relatórios gerenciais completos/exportáveis.", "Trocar 'relatórios' por 'dashboard e histórico filtrável'."],
    ["D14", "Metodologias adotadas", "Afirma uso de StarUML e Figma como parte da construção, mas isso não aparece no código-fonte; o Figma aparece apenas no texto do apêndice.", "Manter como artefatos de projeto, não como evidência de implementação."],
    ["D15", "Metodologias adotadas", "Diz que foram desenvolvidos cadastro de usuários, gerenciamento de salas, agendamento e painel administrativo, omitindo os módulos realmente implementados depois.", "Acrescentar blocos, manutenção, recorrência, avisos, mensagens e histórico."],
    ["D16", "Resultados esperados", "Fala apenas em resultados esperados; o sistema já possui implementação concreta.", "Reescrever como resultados obtidos/funcionalidades alcançadas."],
    ["D17", "Resultados esperados", "Afirma que a gestão poderá acompanhar estatísticas de uso. O código mostra contadores e histórico, mas não estatísticas analíticas de ocupação.", "Limitar a dashboard, contadores e filtros históricos."],
    ["D18", "Conclusão", "Trata notificações automáticas como expansão futura. O código já implementa notificações visuais de mensagens e reservas respondidas.", "Atualizar para mencionar notificações implementadas e reservar como futuro apenas integrações externas ou relatórios avançados."],
    ["D19", "Conclusão", "Afirma integração com sistemas acadêmicos como possibilidade, mas não diferencia o que foi implementado do que é futuro.", "Separar claramente funcionalidades implementadas e oportunidades futuras."],
    ["D20", "Apêndice A", "O texto do diagrama de classes provavelmente não contempla entidades atuais como Aviso, MensagemDireta, campos de manutenção e grupo de recorrência.", "Atualizar o diagrama de classes para incluir User, Bloco, Sala, Reserva, Aviso e MensagemDireta."],
    ["D21", "Apêndice B", "O caso de uso citado no texto é genérico e não contempla aprovar recorrência em lote, enviar mensagens, publicar avisos e gerenciar manutenção.", "Atualizar o diagrama de casos de uso com todos os casos implementados."],
    ["D22", "Apêndice C", "Os protótipos são tratados como interface final, mas o sistema real possui telas adicionais e diferenças visuais.", "Descrever os protótipos como base inicial e anexar capturas reais das telas implementadas."],
    ["D23", "Banco de dados", "O TCC não descreve o modelo implementado. O código possui tabelas users, blocos, salas, reservas, avisos, mensagens_diretas, sessions e tabelas padrão do Laravel.", "Adicionar subseção de modelo de dados com entidades, atributos e relacionamentos."],
    ["D24", "Controle de acesso", "O TCC cita autenticação, mas não explica middleware e redirecionamento por perfil.", "Descrever auth, sessão, perfil admin/professor e middleware CheckAdmin."],
    ["D25", "Regras de negócio", "O TCC não documenta regras de conflito, manutenção, aprovação automática após 12 horas e cancelamento de concorrentes.", "Adicionar subseção de regras de negócio implementadas."],
    ["D26", "Qualidade do texto", "Há inconsistência de tempos verbais e erros como 'serão plicados', 'sera', 'Que aplicados'.", "Revisar para linguagem acadêmica no presente/passado."],
]


CORRECTED_CHAPTERS = [
    ("RESUMO", [
        "Este Trabalho de Conclusão de Curso apresenta o desenvolvimento do UniSalas, um sistema web destinado à organização do processo de solicitação e gerenciamento de reservas de salas acadêmicas. A aplicação foi implementada com Laravel, PHP, Blade, Bootstrap e banco de dados MySQL/MariaDB, contemplando dois perfis principais de acesso: administrador e professor.",
        "O sistema permite que professores solicitem reservas simples ou recorrentes, consultem suas solicitações, acompanhem respostas administrativas, desistam de reservas específicas e troquem mensagens com a administração. Para o administrador, foram implementados recursos de gerenciamento de usuários, blocos, salas, manutenção de espaços, análise de reservas, aprovação ou recusa individual e em lote, histórico filtrável, publicação de avisos e troca de mensagens diretas.",
        "A solução reduz a dependência de controles manuais ao centralizar informações sobre salas, períodos, status das solicitações, comentários e indisponibilidades por manutenção. O sistema também executa regras de negócio para verificar conflitos, cancelar reservas afetadas por manutenção, aprovar solicitações em situações específicas e registrar respostas aos professores. Dessa forma, o UniSalas contribui para maior organização, rastreabilidade e eficiência no uso dos espaços acadêmicos.",
        "Palavras-chave: Reserva de salas; Sistema web; Laravel; Gestão acadêmica; Automação."
    ]),
    ("ABSTRACT", [
        "This final course project presents UniSalas, a web-based system designed to organize classroom reservation requests and management. The application was implemented using Laravel, PHP, Blade, Bootstrap, and a MySQL/MariaDB database, supporting two main access profiles: administrator and teacher.",
        "The system allows teachers to request single or recurring reservations, view their requests, follow administrative responses, cancel specific reservations, and exchange direct messages with the administration. Administrators can manage users, buildings, rooms, maintenance periods, reservation analysis, individual or batch approval and rejection, filtered history, announcements, and direct messages.",
        "The solution reduces dependence on manual controls by centralizing information about rooms, periods, request statuses, comments, and maintenance-related unavailability. It also applies business rules for conflict checking, maintenance-based cancellation, specific automatic approvals, and administrative responses to teachers. Therefore, UniSalas improves organization, traceability, and efficiency in the use of academic spaces.",
        "Keywords: Room reservation; Web system; Laravel; Academic management; Automation."
    ]),
    ("1. INTRODUÇÃO", [
        "A organização e o gerenciamento de salas acadêmicas são atividades essenciais para a continuidade das rotinas de ensino. Quando esse processo é conduzido de forma manual ou dispersa, a instituição fica sujeita a conflitos de horário, falhas de comunicação, dificuldade de acompanhamento e uso inadequado dos espaços físicos.",
        "Diante desse contexto, este trabalho apresenta o UniSalas, um sistema web desenvolvido para centralizar e automatizar partes relevantes do processo de reserva de salas. A aplicação permite que professores registrem solicitações de reserva e acompanhem seus status, enquanto administradores analisam pedidos, gerenciam a infraestrutura cadastrada e mantêm a comunicação com os usuários.",
        "A implementação contempla funcionalidades além do agendamento básico, como cadastro de blocos e salas, observações dos espaços, controle de manutenção por bloco ou sala, reservas recorrentes, aprovação administrativa individual ou em lote, histórico filtrável, avisos institucionais, mensagens diretas e notificações visuais. Dessa forma, o sistema transforma o processo de reserva em um fluxo centralizado, rastreável e mais eficiente.",
        "O projeto foi desenvolvido em arquitetura web utilizando o framework Laravel, banco de dados MySQL/MariaDB e interfaces Blade com Bootstrap. Essa escolha técnica favorece a organização em camadas, a manutenção do código e a evolução futura do sistema."
    ]),
    ("2. DESENVOLVIMENTO", [
        "O desenvolvimento do UniSalas foi orientado pela necessidade de substituir controles manuais por uma aplicação capaz de organizar usuários, salas, reservas e comunicação institucional. A seguir, são descritos os requisitos, a estrutura, a implementação e os resultados obtidos com base no código-fonte do sistema."
    ]),
    ("2.1 LEVANTAMENTO DE REQUISITOS", [
        "A análise do sistema implementado permite identificar requisitos funcionais e não funcionais diretamente refletidos no código-fonte. Entre os requisitos funcionais, destacam-se: autenticação de usuários; diferenciação entre perfil administrador e professor; gerenciamento de usuários; gerenciamento de blocos e salas; definição de manutenção em bloco ou sala; solicitação de reservas simples; solicitação de reservas recorrentes; verificação de disponibilidade; aprovação ou recusa individual de reservas; aprovação ou recusa de grupos recorrentes; cancelamento de conflitos; publicação de avisos; histórico administrativo; acompanhamento de reservas pelo professor; desistência de reserva; troca de mensagens diretas; e controle de notificações de reservas respondidas.",
        "Os requisitos não funcionais observados incluem controle de acesso por middleware, autenticação baseada em sessão, uso de banco de dados relacional, interface responsiva, organização MVC característica do Laravel, rastreabilidade por timestamps, armazenamento de comentários administrativos e limpeza automática de dados antigos em alguns fluxos do sistema.",
        "A funcionalidade de edição de reserva pelo professor não está implementada como fluxo funcional nas rotas do sistema. Existe uma view chamada editar-reserva.blade.php, porém ela não possui rota associada no arquivo routes/web.php e apresenta conteúdo estático. Por isso, a documentação deve tratar edição de reserva como melhoria futura, mantendo como funcionalidade implementada apenas a desistência/cancelamento de uma reserva específica."
    ]),
    ("2.2 ESTRUTURA DO SISTEMA", [
        "O sistema está estruturado em dois perfis principais: administrador e professor. O acesso é controlado por autenticação Laravel e por um middleware administrativo que permite acesso às rotas /admin apenas a usuários com tipo admin ou e-mail administrativo definido no código.",
        "O administrador possui acesso ao dashboard, gerenciamento de usuários, gerenciamento de blocos e salas, manutenção de blocos e salas, fila de aprovação de reservas, histórico geral, avisos e mensagens. Na fila de aprovação, reservas normais e recorrentes são agrupadas para análise, permitindo decisão individual ou decisão em lote para uma solicitação recorrente.",
        "O professor possui acesso ao painel principal, solicitação de nova reserva, minhas reservas, histórico e mensagens. A solicitação pode ser simples ou recorrente; no caso recorrente, o sistema gera datas semanais por até três meses a partir das datas base selecionadas. O professor também pode desistir de uma reserva específica e visualizar comentários administrativos.",
        "A interface é composta por views Blade organizadas em diretórios por domínio: auth, layouts, admin, professor, mensagens e public. O layout principal concentra o menu lateral, os indicadores de mensagens e reservas respondidas e os alertas visuais."
    ]),
    ("2.3 DESENVOLVIMENTO E IMPLEMENTAÇÃO", [
        "O UniSalas foi implementado com PHP 8.2 e Laravel 12. O projeto utiliza rotas web, controllers, models Eloquent, migrations e views Blade. O front-end emprega Bootstrap 5, Font Awesome e SweetAlert2, além de configuração de Vite, Tailwind CSS e Axios no package.json.",
        "A camada de controle é composta principalmente por AuthController e AdminController. O AuthController concentra login e logout, incluindo validação de credenciais, regeneração de sessão e redirecionamento por perfil. O AdminController concentra os fluxos operacionais do sistema, incluindo usuários, blocos, salas, reservas, manutenção, histórico, avisos e mensagens.",
        "Os models principais são User, Bloco, Sala, Reserva, Aviso e MensagemDireta. Os relacionamentos implementados indicam que um usuário possui muitas reservas, um bloco possui muitas salas, uma sala pertence a um bloco, uma sala possui muitas reservas e uma reserva pertence a um usuário e a uma sala.",
        "O banco de dados está configurado para MySQL/MariaDB, com tabelas para usuários, blocos, salas, reservas, avisos, mensagens diretas e estruturas padrão do Laravel, como sessões, cache e jobs. O sistema armazena status, comentários, recorrência, grupo de recorrência e dados de manutenção, permitindo maior rastreabilidade das decisões administrativas."
    ]),
    ("2.4 PROBLEMATIZAÇÃO", [
        "A problemática central do projeto está relacionada à dificuldade de organizar reservas de salas quando informações de disponibilidade, solicitações, respostas e manutenção estão dispersas. Essa condição pode gerar conflitos de horários, falta de transparência, atrasos e retrabalho para professores e administradores.",
        "O UniSalas responde a esse problema por meio de um fluxo digital centralizado. O professor registra a solicitação e o sistema verifica a existência de conflito por sala, data e período. Quando a sala ou o bloco está em manutenção, a solicitação recebe tratamento específico e o professor visualiza o motivo da indisponibilidade. Quando existem conflitos em reservas recorrentes, o sistema coloca datas em análise, permitindo decisão administrativa.",
        "Diferentemente de uma confirmação automática universal, o sistema implementado trabalha com status e regras: pendente, em análise, aprovada, rejeitada e cancelada. Também há aprovação automática em situações específicas, como solicitações realizadas no mesmo dia ou pendências antigas processadas pela rotina administrativa. Essa abordagem preserva controle administrativo e reduz sobreposições indevidas."
    ]),
    ("Tabela 1 - Soluções implementadas", [
        "Falta de controle centralizado: o sistema centraliza usuários, blocos, salas, reservas, avisos e mensagens em banco de dados.",
        "Conflitos de horário: a aplicação verifica sala, data e período antes de registrar o status inicial da reserva.",
        "Indisponibilidade por manutenção: administradores podem colocar bloco ou sala em manutenção, cancelando reservas afetadas e registrando motivo.",
        "Solicitações recorrentes: professores podem gerar reservas semanais por até três meses, agrupadas por grupo de recorrência.",
        "Análise administrativa: administradores podem aprovar ou recusar reservas individualmente e grupos recorrentes em lote.",
        "Falta de comunicação: o sistema possui avisos e mensagens diretas entre administradores e professores.",
        "Acompanhamento: o professor possui tela de minhas reservas e histórico, enquanto o administrador possui histórico geral filtrável."
    ]),
    ("2.5 METODOLOGIAS ADOTADAS", [
        "O desenvolvimento seguiu uma abordagem incremental, com evolução progressiva das funcionalidades do sistema. A organização por módulos permitiu implementar primeiro a autenticação e os cadastros básicos, avançando posteriormente para reservas, recorrência, manutenção, histórico, avisos, mensagens e notificações.",
        "Foram aplicados princípios compatíveis com metodologias ágeis, como simplicidade, validação frequente, refatoração e ajustes por necessidade do usuário. A implementação também recebeu correções incrementais, como melhoria de contraste na interface de manutenção, aprovação em lote de reservas recorrentes e ajuste no comportamento de notificações de reservas respondidas.",
        "Na perspectiva técnica, o sistema adota arquitetura MVC por meio do Laravel. As rotas definem os pontos de entrada, os controllers coordenam regras de negócio, os models representam entidades persistentes e as views Blade compõem a interface. O banco de dados relacional sustenta a persistência das informações de usuários, salas e reservas."
    ]),
    ("2.6 RESULTADOS OBTIDOS E BENEFÍCIOS", [
        "Como resultado, foi desenvolvido um sistema funcional para apoiar a gestão de reservas acadêmicas. A aplicação permite organizar a infraestrutura em blocos e salas, controlar manutenção, cadastrar usuários, receber solicitações, analisar reservas e manter histórico das decisões.",
        "Para professores, os principais benefícios são a solicitação digital de reservas, a possibilidade de recorrência, o acompanhamento de status, o recebimento de respostas administrativas e a comunicação por mensagens. Para administradores, destacam-se a centralização de solicitações, a análise de conflitos, a aprovação em lote de recorrências, os avisos institucionais e o histórico filtrável.",
        "O sistema não possui, no estado atual, relatórios gerenciais analíticos completos ou exportáveis. Entretanto, apresenta dashboard com contadores, histórico por filtros e registros suficientes para acompanhamento operacional. Esses recursos podem servir como base para futuras estatísticas de ocupação, indicadores de uso e relatórios administrativos."
    ]),
    ("3. CONCLUSÃO", [
        "O desenvolvimento do UniSalas resultou em uma aplicação web capaz de organizar o processo de reserva de salas acadêmicas, substituindo parte significativa do controle manual por um fluxo digital centralizado. O sistema implementado oferece recursos para professores e administradores, contemplando autenticação, gerenciamento de infraestrutura, solicitação de reservas, aprovação administrativa, histórico, avisos, mensagens e controle de manutenção.",
        "A solução contribui para a redução de conflitos de agendamento ao verificar sala, data e período, além de registrar status e comentários administrativos. O tratamento de reservas recorrentes, a decisão em lote e o cancelamento de reservas afetadas por manutenção ampliam a aderência do sistema às necessidades reais de uso institucional.",
        "Embora o projeto já apresente uma base funcional consistente, há oportunidades de melhoria, como separação do controller principal em controllers especializados, criação de relatórios gerenciais, padronização formal dos status de reserva, fortalecimento das constraints de banco e ampliação dos testes automatizados. Ainda assim, o sistema desenvolvido demonstra viabilidade técnica e utilidade prática para a organização de salas acadêmicas."
    ]),
    ("APÊNDICES - AJUSTES RECOMENDADOS", [
        "O Diagrama de Classes deve ser atualizado para conter as entidades User, Bloco, Sala, Reserva, Aviso e MensagemDireta, além dos atributos de manutenção em blocos e salas e os campos de recorrência em reservas.",
        "O Diagrama de Casos de Uso deve incluir, além de agendar sala e gerenciar usuários, os casos: gerenciar blocos, gerenciar salas, controlar manutenção, solicitar reserva recorrente, aprovar grupo recorrente, publicar aviso, enviar mensagem direta, consultar histórico e desistir de reserva.",
        "Os protótipos devem ser apresentados como base visual inicial. Para representar fielmente o sistema implementado, recomenda-se anexar capturas reais das telas atuais: dashboard administrativo, usuários, blocos e salas, aprovação de reservas, histórico geral, painel do professor, nova reserva, minhas reservas e mensagens."
    ]),
]


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.bold = header
    if header:
        r.font.color.rgb = RGBColor.from_string("0B2545")
        shade(cell, "F2F4F7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, data, widths):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            set_text(table.cell(i, j), value, i == 0)
            table.cell(i, j).width = Inches(widths[j])
    doc.add_paragraph()


def build_docx():
    doc = Document()
    setup_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Auditoria entre TCC e Código-fonte")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"UniSalas - divergências e capítulos corrigidos\nData: {TODAY}").italic = True

    doc.add_heading("1. Escopo da auditoria", level=1)
    doc.add_paragraph(
        "Esta auditoria compara o conteúdo textual extraído do arquivo tcc_luid.pdf com o código-fonte do projeto Laravel localizado em c:\\xampp\\htdocs\\reserva-salas. "
        "Foram analisados rotas, controllers, models, migrations, views Blade, arquivos de dependências e a documentação técnica previamente gerada."
    )
    doc.add_paragraph(
        "A análise concentra-se na fidelidade entre o que o TCC afirma e o que o sistema realmente implementa. Não foram avaliados os conteúdos visuais internos dos diagramas extraídos como imagem, apenas as descrições textuais dos apêndices e as evidências disponíveis no código."
    )

    doc.add_heading("2. Divergências encontradas", level=1)
    add_table(doc, [["ID", "Local", "Divergência", "Correção recomendada"]] + DIVERGENCES, [0.45, 1.1, 2.8, 2.15])

    doc.add_page_break()
    doc.add_heading("3. Versão corrigida dos capítulos", level=1)
    doc.add_paragraph(
        "Os textos a seguir foram reescritos para representar fielmente o sistema implementado no código-fonte. Eles podem substituir ou complementar os capítulos equivalentes do TCC."
    )
    for heading, paragraphs in CORRECTED_CHAPTERS:
        level = 1 if heading in {"RESUMO", "ABSTRACT"} or heading.startswith(("1.", "2.", "3.", "APÊNDICES")) else 2
        doc.add_heading(heading, level=level)
        if heading == "Tabela 1 - Soluções implementadas":
            for item in paragraphs:
                doc.add_paragraph(item, style="List Bullet")
        else:
            for text in paragraphs:
                doc.add_paragraph(text)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Auditoria TCC x Código-fonte - UniSalas").font.size = Pt(9)
    OUT_DIR.mkdir(exist_ok=True)
    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleX", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=10))
    styles.add(ParagraphStyle("SubX", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78"), spaceAfter=16))
    styles.add(ParagraphStyle("H1X", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor("#2E74B5"), spaceBefore=11, spaceAfter=7))
    styles.add(ParagraphStyle("H2X", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#2E74B5"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle("BodyX", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12.4, alignment=TA_JUSTIFY, spaceAfter=6))
    styles.add(ParagraphStyle("BulletX", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12, leftIndent=14, firstLineIndent=-8, spaceAfter=4))
    styles.add(ParagraphStyle("CellX", parent=styles["BodyText"], fontName="Helvetica", fontSize=6.6, leading=8.2, alignment=TA_LEFT))
    styles.add(ParagraphStyle("CellHeadX", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=6.8, leading=8.4, textColor=colors.HexColor("#0B2545"), alignment=TA_LEFT))
    return styles


def para(text, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def add_pdf_table(story, styles, data, widths):
    rows = []
    for i, row in enumerate(data):
        style = styles["CellHeadX"] if i == 0 else styles["CellX"]
        rows.append([para(str(cell), style) for cell in row])
    table = Table(rows, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C2CC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch)
    story = [
        para("Auditoria entre TCC e Código-fonte", styles["TitleX"]),
        para(f"UniSalas - divergências e capítulos corrigidos<br/>Data: {TODAY}", styles["SubX"]),
        para("1. Escopo da auditoria", styles["H1X"]),
        para("Esta auditoria compara o conteúdo textual extraído do arquivo tcc_luid.pdf com o código-fonte do projeto Laravel localizado em c:\\xampp\\htdocs\\reserva-salas. Foram analisados rotas, controllers, models, migrations, views Blade, arquivos de dependências e a documentação técnica previamente gerada.", styles["BodyX"]),
        para("A análise concentra-se na fidelidade entre o que o TCC afirma e o que o sistema realmente implementa. Não foram avaliados os conteúdos visuais internos dos diagramas extraídos como imagem, apenas as descrições textuais dos apêndices e as evidências disponíveis no código.", styles["BodyX"]),
        para("2. Divergências encontradas", styles["H1X"]),
    ]
    add_pdf_table(story, styles, [["ID", "Local", "Divergência", "Correção recomendada"]] + DIVERGENCES, [0.42 * inch, 1.03 * inch, 2.78 * inch, 2.27 * inch])
    story.append(PageBreak())
    story.append(para("3. Versão corrigida dos capítulos", styles["H1X"]))
    story.append(para("Os textos a seguir foram reescritos para representar fielmente o sistema implementado no código-fonte. Eles podem substituir ou complementar os capítulos equivalentes do TCC.", styles["BodyX"]))
    for heading, paragraphs in CORRECTED_CHAPTERS:
        style = styles["H1X"] if heading in {"RESUMO", "ABSTRACT"} or heading.startswith(("1.", "2.", "3.", "APÊNDICES")) else styles["H2X"]
        story.append(para(heading, style))
        if heading == "Tabela 1 - Soluções implementadas":
            for item in paragraphs:
                story.append(para(f"• {item}", styles["BulletX"]))
        else:
            for text in paragraphs:
                story.append(para(text, styles["BodyX"]))
    doc.build(story)


if __name__ == "__main__":
    OUT_DIR.mkdir(exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
