from pathlib import Path
import zipfile

import fitz
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Documentacao_Tecnica_Reserva_Salas.docx"
PDF = ROOT / "docs" / "Documentacao_Tecnica_Reserva_Salas.pdf"
PNG_DIR = ROOT / "docs" / "pdf_preview"


def main():
    doc = Document(DOCX)
    with zipfile.ZipFile(DOCX) as zf:
        required = {
            "[Content_Types].xml",
            "word/document.xml",
            "word/styles.xml",
        }
        missing = sorted(required - set(zf.namelist()))
        if missing:
            raise RuntimeError(f"DOCX sem partes obrigatórias: {missing}")

    pdf = fitz.open(PDF)
    PNG_DIR.mkdir(exist_ok=True)
    for page_index in range(len(pdf)):
        page = pdf.load_page(page_index)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
        pix.save(PNG_DIR / f"page-{page_index + 1}.png")

    print(f"DOCX bytes: {DOCX.stat().st_size}")
    print(f"DOCX paragraphs: {len(doc.paragraphs)}")
    print(f"DOCX tables: {len(doc.tables)}")
    print(f"PDF bytes: {PDF.stat().st_size}")
    print(f"PDF pages: {len(pdf)}")
    print(f"PDF preview dir: {PNG_DIR}")


if __name__ == "__main__":
    main()
