from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "Documentacao_Tecnica_Reserva_Salas.docx"
PDF_PATH = OUT_DIR / "Documentacao_Tecnica_Reserva_Salas.pdf"


TITLE = "Documentação Técnica do Sistema de Reserva de Salas"
SUBTITLE = "Projeto Laravel: reserva-salas"
TODAY = date.today().strftime("%d/%m/%Y")


TECHNOLOGIES = [
    ["Camada", "Tecnologia", "Uso no projeto"],
    ["Back-end", "PHP 8.2 e Laravel 12", "Rotas, controllers, models Eloquent, autenticação, sessões, migrations e validações."],
    ["Front-end", "Blade, Bootstrap 5, Font Awesome e SweetAlert2", "Renderização das telas, layout responsivo, ícones, alertas e feedback visual."],
    ["Build front-end", "Vite, Tailwind CSS, Axios", "Configuração disponível no package.json para build e desenvolvimento."],
    ["Banco de dados", "MySQL/MariaDB em 127.0.0.1:3306", "Persistência configurada para a base reserva_salas."],
    ["Testes", "PHPUnit 11", "Suíte básica de testes automatizados do Laravel."],
    ["Ambiente", "XAMPP / Windows", "Projeto localizado em c:\\xampp\\htdocs\\reserva-salas."],
]


FUNCTIONAL_REQUIREMENTS = [
    ["ID", "Requisito funcional", "Implementação observada"],
    ["RF01", "Autenticar usuários por e-mail e senha.", "AuthController valida credenciais, regenera sessão e redireciona por perfil."],
    ["RF02", "Separar acesso de administrador e professor.", "Middleware CheckAdmin restringe rotas /admin a usuários tipo admin ou e-mail administrativo."],
    ["RF03", "Gerenciar usuários.", "Administrador lista, cria, atualiza e exclui usuários; exclusão remove reservas e mensagens vinculadas."],
    ["RF04", "Gerenciar blocos.", "Administrador cadastra, edita, exclui e define cor visual do bloco."],
    ["RF05", "Gerenciar salas por bloco.", "Administrador cadastra, edita, exclui e consulta salas associadas aos blocos."],
    ["RF06", "Controlar manutenção de blocos e salas.", "Sistema registra manutenção ativa, fim, indeterminação e aviso, cancelando reservas futuras afetadas."],
    ["RF07", "Solicitar reserva simples.", "Professor seleciona sala, data e período; sistema valida conflito e define status inicial."],
    ["RF08", "Solicitar reserva recorrente.", "Professor informa datas base; sistema gera ocorrências semanais por até três meses."],
    ["RF09", "Verificar disponibilidade antes da reserva.", "Endpoint /professor/verificar retorna disponibilidade e mensagem de conflito/manutenção."],
    ["RF10", "Aprovar ou recusar reservas individualmente.", "Administrador altera status e comentário administrativo por solicitação."],
    ["RF11", "Aprovar ou recusar grupo recorrente.", "Administrador pode decidir todas as reservas pendentes de um grupo de recorrência em lote."],
    ["RF12", "Cancelar conflitos automaticamente.", "Ao aprovar uma reserva, solicitações concorrentes pendentes/em análise são canceladas."],
    ["RF13", "Aprovar pendentes expiradas.", "Reservas pendentes há 12 horas são aprovadas automaticamente quando o admin acessa a fila."],
    ["RF14", "Exibir histórico administrativo.", "Administrador filtra histórico por data, bloco, sala e período."],
    ["RF15", "Exibir painel e reservas do professor.", "Professor vê totais, lembretes, reservas ativas, histórico e respostas administrativas."],
    ["RF16", "Permitir desistência de reserva.", "Professor pode excluir uma reserva específica vinculada ao seu usuário."],
    ["RF17", "Publicar e remover avisos.", "Administrador publica avisos exibidos no dashboard e no painel do professor."],
    ["RF18", "Trocar mensagens diretas.", "Usuários enviam mensagens, visualizam conversas, marcam mensagens como lidas e apagam chats."],
    ["RF19", "Controlar notificações de reservas respondidas.", "Ao entrar em Minhas Reservas, a sessão marca a visualização e encerra a notificação recorrente."],
]


NON_FUNCTIONAL_REQUIREMENTS = [
    ["ID", "Requisito não funcional", "Evidência no projeto"],
    ["RNF01", "Segurança de sessão", "Uso de Auth, regenerate(), invalidate() e regenerateToken() no login/logout."],
    ["RNF02", "Controle de acesso", "Rotas administrativas protegidas por middleware e rotas gerais protegidas por auth."],
    ["RNF03", "Integridade relacional", "Relacionamentos Eloquent e chaves estrangeiras entre usuários, salas, blocos e reservas."],
    ["RNF04", "Usabilidade", "Interface responsiva com sidebar, badges, cards, colapsos e feedback por toast."],
    ["RNF05", "Manutenibilidade", "Estrutura MVC Laravel com views agrupadas por domínio: admin, professor, mensagens, auth e public."],
    ["RNF06", "Disponibilidade operacional", "Limpeza automática de mensagens e reservas antigas reduz crescimento indefinido dos dados."],
    ["RNF07", "Consistência de regras", "Conflitos, manutenção, recorrência e aprovação em lote centralizados em métodos do controller."],
    ["RNF08", "Rastreabilidade", "Reservas armazenam comentário do professor, comentário do administrador, status e timestamps."],
    ["RNF09", "Portabilidade local", "Aplicação segue padrão Laravel e pode ser executada via PHP/Composer/Vite em ambiente XAMPP."],
]


USE_CASES = [
    ["ID", "Caso de uso", "Ator principal", "Descrição"],
    ["UC01", "Realizar login", "Usuário", "Informar e-mail e senha para acessar o painel correspondente ao perfil."],
    ["UC02", "Gerenciar usuários", "Administrador", "Criar, listar, editar dados e excluir usuários do sistema."],
    ["UC03", "Gerenciar blocos e salas", "Administrador", "Manter a estrutura física disponível para reservas."],
    ["UC04", "Colocar bloco/sala em manutenção", "Administrador", "Definir manutenção por prazo determinado ou indeterminado e cancelar reservas afetadas."],
    ["UC05", "Solicitar reserva", "Professor", "Escolher bloco, sala, período e data para abrir solicitação."],
    ["UC06", "Solicitar reserva recorrente", "Professor", "Escolher datas base e gerar ocorrências semanais por três meses."],
    ["UC07", "Analisar solicitações", "Administrador", "Aprovar ou recusar reservas pendentes ou em análise."],
    ["UC08", "Analisar recorrência em lote", "Administrador", "Aprovar ou recusar todas as reservas de uma solicitação recorrente."],
    ["UC09", "Consultar histórico", "Administrador/Professor", "Visualizar registros de reservas com status e comentários."],
    ["UC10", "Enviar mensagem direta", "Administrador/Professor", "Trocar mensagens entre perfis, acompanhar mensagens não lidas e apagar conversas."],
    ["UC11", "Publicar aviso", "Administrador", "Criar comunicados institucionais para exibição nos painéis."],
    ["UC12", "Visualizar resposta de reserva", "Professor", "Entrar em Minhas Reservas e encerrar a notificação de resposta já visualizada."],
]


DATABASE_MODEL = [
    ["Tabela", "Campos principais", "Relacionamentos e observações"],
    ["users", "id, name, email, tipo, telefone, is_admin, password, timestamps", "Usuário possui muitas reservas. Campo tipo separa admin/professor."],
    ["blocos", "id, nome, cor, manutencao_ativa, manutencao_fim, manutencao_indeterminada, manutencao_aviso, timestamps", "Bloco possui muitas salas. Manutenção em bloco afeta todas as salas vinculadas."],
    ["salas", "id, nome, capacidade, observacao, bloco_id, campos de manutenção, timestamps", "Sala pertence a bloco e possui muitas reservas. Exclusão do bloco remove salas em cascata."],
    ["reservas", "id, user_id, sala_id, data_reserva, periodo, status, recorrente, grupo_recorrencia, observacao, comentario_professor, comentario_adm, timestamps", "Reserva pertence a usuário e sala. grupo_recorrencia agrupa solicitações semanais."],
    ["avisos", "id, titulo, mensagem, timestamps", "Avisos exibidos nos painéis e administrados pelo admin."],
    ["mensagens_diretas", "id, remetente_id, destinatario_id, mensagem, lida, timestamps", "Mensagens entre usuários; lida controla notificações."],
    ["sessions", "id, user_id, ip_address, user_agent, payload, last_activity", "Persistência de sessão Laravel."],
    ["password_reset_tokens", "email, token, created_at", "Tabela padrão para recuperação de senha."],
]


SCREENS = [
    ["Área", "Tela", "Arquivo Blade", "Descrição"],
    ["Autenticação", "Login", "resources/views/auth/login.blade.php", "Entrada do usuário com credenciais."],
    ["Layout", "Base autenticada", "resources/views/layouts/app.blade.php", "Sidebar, alertas, notificações e estrutura comum."],
    ["Admin", "Dashboard", "resources/views/admin/dashboard.blade.php", "Métricas, avisos e visão geral."],
    ["Admin", "Usuários", "resources/views/admin/usuarios/index.blade.php", "Lista compacta, edição, exclusão e reservas vinculadas."],
    ["Admin", "Criar usuário", "resources/views/admin/usuarios/create.blade.php", "Cadastro de novo acesso."],
    ["Admin", "Blocos e salas", "resources/views/admin/blocos-salas/index.blade.php", "Gestão da infraestrutura e manutenção."],
    ["Admin", "Criar/editar bloco", "resources/views/admin/blocos-salas/blocos-create.blade.php; blocos-edit.blade.php", "Cadastro e atualização de blocos."],
    ["Admin", "Criar/editar sala", "resources/views/admin/blocos-salas/salas-create.blade.php; salas-edit.blade.php", "Cadastro e atualização de salas."],
    ["Admin", "Aprovar reservas", "resources/views/admin/reservas/aprovar.blade.php", "Fila de solicitações normais e recorrentes."],
    ["Admin", "Histórico geral", "resources/views/admin/historico.blade.php", "Consulta de reservas finalizadas com filtros."],
    ["Professor", "Painel", "resources/views/professor/painel.blade.php", "Resumo de reservas, lembretes, avisos e mensagem ao admin."],
    ["Professor", "Nova reserva", "resources/views/professor/nova-reserva.blade.php", "Solicitação simples ou recorrente com verificação de disponibilidade."],
    ["Professor", "Minhas reservas", "resources/views/professor/minhas-reservas.blade.php", "Reservas ativas e respostas da administração."],
    ["Professor", "Histórico", "resources/views/professor/historico.blade.php", "Consulta das solicitações do professor."],
    ["Mensagens", "Lista de contatos", "resources/views/mensagens/index.blade.php", "Envio de nova mensagem e contatos com mensagens não lidas."],
    ["Mensagens", "Chat", "resources/views/mensagens/chat.blade.php", "Conversa direta entre usuários."],
    ["Público", "Home/mapa/welcome", "resources/views/public/*.blade.php", "Telas públicas existentes, embora a rota raiz redirecione para login."],
]


ROUTES = [
    ["Grupo", "Rotas", "Finalidade"],
    ["Público", "GET /, GET/POST /login, GET /logout", "Entrada, autenticação e encerramento de sessão."],
    ["Mensagens", "GET /mensagens, GET /chat/{id}, POST /mensagem/enviar, DELETE /chat/{id}", "Listagem de contatos, conversa, envio e exclusão de chat."],
    ["Admin", "GET /admin/dashboard", "Painel administrativo."],
    ["Admin usuários", "/admin/usuarios, /admin/usuarios/novo, salvar, atualizar, excluir", "CRUD de usuários."],
    ["Admin infraestrutura", "/admin/blocos, /admin/sala/*, rotas de manutenção", "CRUD de blocos/salas e manutenção."],
    ["Admin reservas", "/admin/reservas, status individual, status por grupo, cancelar, histórico", "Análise e histórico de reservas."],
    ["Admin avisos", "/admin/avisos/salvar, /admin/avisos/excluir/{id}", "Criação e remoção de avisos."],
    ["Professor", "/professor/painel, minhas-reservas, solicitar, histórico, desistir", "Fluxo operacional do professor."],
    ["Professor API", "/professor/api/salas/{bloco_id}, /professor/verificar", "Carregamento de salas e verificação de disponibilidade."],
]


ARCHITECTURE_ITEMS = [
    "A arquitetura segue o padrão MVC do Laravel. As rotas em routes/web.php apontam para AuthController e AdminController; os controllers coordenam validação, regras de negócio e escolha das views.",
    "A camada de modelo utiliza Eloquent para representar User, Bloco, Sala, Reserva, Aviso e MensagemDireta. Os relacionamentos principais são User 1:N Reserva, Bloco 1:N Sala e Sala 1:N Reserva.",
    "A interface é implementada com Blade, Bootstrap e CSS embutido no layout principal. O layout centraliza menu lateral, feedbacks SweetAlert2 e indicadores de notificações.",
    "A persistência utiliza migrations Laravel em MySQL/MariaDB. As migrations incluem proteção condicional com Schema::hasColumn/hasTable em extensões recentes do modelo.",
    "O controle de acesso é dividido em middleware auth para usuários autenticados e middleware admin para a área administrativa.",
    "As regras críticas de reserva ficam centralizadas no AdminController: detecção de conflitos, manutenção, recorrência, aprovação automática, cancelamento de concorrentes e atualização em lote.",
]


IMPLEMENTED_FEATURES = [
    "Autenticação por sessão com redirecionamento para painel administrativo ou painel do professor.",
    "Dashboard administrativo com contadores, avisos e mensagens não lidas.",
    "CRUD de usuários, com preservação/remoção controlada de dados relacionados.",
    "CRUD de blocos e salas, incluindo cor do bloco e observações de sala.",
    "Gestão de manutenção em bloco ou sala, com opção de prazo indeterminado e cancelamento de reservas afetadas.",
    "Solicitação de reserva simples com validação de data futura e disponibilidade.",
    "Solicitação recorrente semanal por até três meses, agrupada por UUID.",
    "Fila administrativa de aprovação com agrupamento de recorrências, análise individual e decisão em lote.",
    "Regras de conflito que cancelam solicitações concorrentes após aprovação.",
    "Aprovação automática de pendências antigas quando a fila administrativa é acessada.",
    "Histórico administrativo filtrável por data, bloco, sala e período.",
    "Painel do professor com métricas, lembrete do dia, avisos e envio de mensagem.",
    "Minhas reservas e histórico do professor com status e comentários administrativos.",
    "Mensagens diretas entre administradores e professores, com indicador de não lidas.",
    "Controle de notificação de reservas respondidas por marcador de visualização em sessão.",
    "Limpeza automática de mensagens e reservas anteriores a seis meses.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_docx_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_docx_table(doc, data, widths=None):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths is None:
        widths = [6.5 / len(data[0])] * len(data[0])
    for row_idx, row in enumerate(data):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, str(value), bold=row_idx == 0, color="0B2545" if row_idx == 0 else "000000")
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            cell.width = Inches(widths[col_idx])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx():
    doc = Document()
    configure_docx_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(SUBTITLE)
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Data de emissão: {TODAY}\nBase analisada: c:\\xampp\\htdocs\\reserva-salas").italic = True

    doc.add_paragraph()
    add_docx_table(doc, [["Item", "Descrição"], ["Tipo de documento", "Relatório técnico em formato acadêmico"], ["Escopo", "Análise do projeto Laravel, código-fonte, rotas, models, migrations e views"], ["Preset visual", "standard_business_brief"]], [1.7, 4.8])
    doc.add_page_break()

    doc.add_heading("1. Introdução", level=1)
    doc.add_paragraph(
        "Este documento apresenta a análise técnica do sistema Laravel denominado reserva-salas. "
        "O projeto tem como finalidade apoiar a reserva de salas em ambiente institucional, com perfis de administrador e professor, "
        "controle de infraestrutura, fluxo de aprovação, comunicação por mensagens diretas e registro histórico das solicitações."
    )
    doc.add_paragraph(
        "A documentação foi produzida a partir da leitura dos arquivos do projeto, incluindo rotas, controllers, models, migrations, "
        "views Blade e arquivos de dependências. O conteúdo descreve o estado implementado do sistema e suas principais regras de negócio."
    )

    doc.add_heading("2. Objetivo", level=1)
    doc.add_paragraph(
        "O objetivo geral do sistema é organizar o processo de solicitação, análise e acompanhamento de reservas de salas, "
        "reduzindo conflitos de agenda e centralizando a gestão de blocos, salas, manutenção, avisos e comunicação entre professores e administradores."
    )
    doc.add_paragraph("Objetivos específicos:")
    add_bullets(doc, [
        "Permitir que professores solicitem reservas simples ou recorrentes.",
        "Permitir que administradores aprovem, recusem, cancelem e auditem solicitações.",
        "Evitar conflitos de sala, data e período por meio de regras automáticas.",
        "Registrar respostas e comentários administrativos para rastreabilidade.",
        "Manter professores informados por avisos, mensagens e notificações de reserva respondida.",
    ])

    doc.add_heading("3. Requisitos funcionais", level=1)
    add_docx_table(doc, FUNCTIONAL_REQUIREMENTS, [0.55, 2.1, 3.85])

    doc.add_heading("4. Requisitos não funcionais", level=1)
    add_docx_table(doc, NON_FUNCTIONAL_REQUIREMENTS, [0.65, 1.9, 3.95])

    doc.add_heading("5. Casos de uso", level=1)
    add_docx_table(doc, USE_CASES, [0.55, 1.65, 1.25, 4.05])

    doc.add_heading("6. Arquitetura do sistema", level=1)
    add_bullets(doc, ARCHITECTURE_ITEMS)
    doc.add_heading("6.1 Mapa de rotas", level=2)
    add_docx_table(doc, ROUTES, [1.25, 2.55, 2.7])

    doc.add_heading("7. Modelo de banco de dados", level=1)
    doc.add_paragraph(
        "O banco está configurado para MySQL/MariaDB, com conexão mysql, host 127.0.0.1, porta 3306 e database reserva_salas. "
        "Credenciais sensíveis não foram reproduzidas neste documento."
    )
    add_docx_table(doc, DATABASE_MODEL, [1.25, 2.95, 2.3])
    doc.add_paragraph(
        "Observação técnica: existem migrations duplicadas ou evolutivas para mensagens diretas e cor de blocos. "
        "Algumas migrations usam verificações Schema::hasTable e Schema::hasColumn para evitar erro quando a estrutura já existe."
    )

    doc.add_heading("8. Telas implementadas", level=1)
    add_docx_table(doc, SCREENS, [0.95, 1.25, 2.25, 2.05])

    doc.add_heading("9. Funcionalidades implementadas", level=1)
    add_bullets(doc, IMPLEMENTED_FEATURES)

    doc.add_heading("10. Tecnologias utilizadas", level=1)
    add_docx_table(doc, TECHNOLOGIES, [1.15, 1.75, 3.6])

    doc.add_heading("11. Fluxos principais", level=1)
    doc.add_heading("11.1 Fluxo de reserva simples", level=2)
    add_bullets(doc, [
        "Professor acessa Solicitar Reserva.",
        "Seleciona bloco, sala, data e período.",
        "Sistema consulta disponibilidade por sala, data e período.",
        "Se houver manutenção, a reserva é cancelada com motivo administrativo.",
        "Se houver conflito em solicitação simples, a reserva é cancelada por ocupação.",
        "Se a data for o mesmo dia, a reserva é aprovada automaticamente; caso contrário, fica pendente.",
    ])
    doc.add_heading("11.2 Fluxo de reserva recorrente", level=2)
    add_bullets(doc, [
        "Professor ativa o modo recorrente e escolhe datas base.",
        "Sistema gera datas semanais até três meses a partir de cada data base.",
        "Todas as reservas recebem o mesmo grupo_recorrencia.",
        "Datas com conflito entram em análise, e datas sem conflito seguem pendentes ou aprovadas conforme regra.",
        "Administrador pode analisar cada data individualmente ou aprovar/recusar todas no fim da lista do grupo.",
    ])
    doc.add_heading("11.3 Fluxo de manutenção", level=2)
    add_bullets(doc, [
        "Administrador ativa manutenção em bloco ou sala.",
        "Define prazo final ou tempo indeterminado e aviso opcional.",
        "Sistema cancela reservas futuras pendentes, em análise ou aprovadas dentro do período afetado.",
        "Novas solicitações para salas em manutenção recebem mensagem de indisponibilidade.",
    ])

    doc.add_heading("12. Considerações técnicas e oportunidades de melhoria", level=1)
    add_bullets(doc, [
        "Separar responsabilidades do AdminController em controllers específicos para usuários, infraestrutura, reservas, avisos e mensagens.",
        "Criar Form Requests para centralizar validações e mensagens de erro.",
        "Adicionar foreign keys explícitas para remetente_id e destinatario_id em mensagens_diretas.",
        "Normalizar status de reserva para enum ou constantes de domínio, evitando variações como recusada e rejeitada.",
        "Expandir testes automatizados para autenticação, reserva recorrente, manutenção, conflitos e notificações.",
        "Revisar migrations duplicadas para reduzir risco de inconsistência em ambientes novos.",
        "Adicionar migrations ou jobs de limpeza dedicados em vez de acionar limpeza apenas ao acessar telas.",
    ])

    doc.add_heading("13. Conclusão", level=1)
    doc.add_paragraph(
        "O sistema reserva-salas apresenta uma base funcional completa para gestão institucional de reservas, contemplando autenticação, "
        "perfis, infraestrutura, solicitações simples e recorrentes, aprovação administrativa, manutenção, avisos, mensagens e histórico. "
        "A aplicação utiliza padrões reconhecidos do Laravel e organiza a interface em telas Blade especializadas por área."
    )
    doc.add_paragraph(
        "Do ponto de vista técnico, o projeto já cobre as principais necessidades operacionais do domínio. "
        "As melhorias recomendadas concentram-se em modularização do controller principal, fortalecimento das constraints de banco, "
        "padronização de status e ampliação da cobertura de testes, medidas que aumentariam a manutenção e a confiabilidade em produção."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Documentação técnica - Sistema de Reserva de Salas").font.size = Pt(9)

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleCustom", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=12))
    styles.add(ParagraphStyle("SubtitleCustom", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78"), spaceAfter=18))
    styles.add(ParagraphStyle("H1Custom", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7))
    styles.add(ParagraphStyle("H2Custom", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#2E74B5"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle("BodyCustom", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12.5, alignment=TA_JUSTIFY, spaceAfter=6))
    styles.add(ParagraphStyle("BulletCustom", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.3, leading=12, leftIndent=14, firstLineIndent=-8, spaceAfter=4))
    styles.add(ParagraphStyle("CellCustom", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9, alignment=TA_LEFT))
    styles.add(ParagraphStyle("CellHeader", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.4, leading=9, textColor=colors.HexColor("#0B2545"), alignment=TA_LEFT))
    return styles


def p(text, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def add_pdf_table(story, styles, data, col_widths):
    rows = []
    for idx, row in enumerate(data):
        style = styles["CellHeader"] if idx == 0 else styles["CellCustom"]
        rows.append([p(str(cell), style) for cell in row])
    table = Table(rows, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C2CC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))


def add_pdf_bullets(story, styles, items):
    for item in items:
        story.append(p(f"• {item}", styles["BulletCustom"]))


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    story = []

    story.append(p(TITLE, styles["TitleCustom"]))
    story.append(p(f"{SUBTITLE}<br/>Data de emissão: {TODAY}<br/>Base analisada: c:\\xampp\\htdocs\\reserva-salas", styles["SubtitleCustom"]))
    add_pdf_table(story, styles, [["Item", "Descrição"], ["Tipo de documento", "Relatório técnico em formato acadêmico"], ["Escopo", "Análise do projeto Laravel, código-fonte, rotas, models, migrations e views"], ["Preset visual", "standard_business_brief"]], [1.7 * inch, 4.8 * inch])
    story.append(PageBreak())

    sections = [
        ("1. Introdução", [
            "Este documento apresenta a análise técnica do sistema Laravel denominado reserva-salas. O projeto tem como finalidade apoiar a reserva de salas em ambiente institucional, com perfis de administrador e professor, controle de infraestrutura, fluxo de aprovação, comunicação por mensagens diretas e registro histórico das solicitações.",
            "A documentação foi produzida a partir da leitura dos arquivos do projeto, incluindo rotas, controllers, models, migrations, views Blade e arquivos de dependências. O conteúdo descreve o estado implementado do sistema e suas principais regras de negócio.",
        ]),
        ("2. Objetivo", [
            "O objetivo geral do sistema é organizar o processo de solicitação, análise e acompanhamento de reservas de salas, reduzindo conflitos de agenda e centralizando a gestão de blocos, salas, manutenção, avisos e comunicação entre professores e administradores.",
        ]),
    ]
    for heading, paragraphs in sections:
        story.append(p(heading, styles["H1Custom"]))
        for text in paragraphs:
            story.append(p(text, styles["BodyCustom"]))

    story.append(p("Objetivos específicos", styles["H2Custom"]))
    add_pdf_bullets(story, styles, [
        "Permitir que professores solicitem reservas simples ou recorrentes.",
        "Permitir que administradores aprovem, recusem, cancelem e auditem solicitações.",
        "Evitar conflitos de sala, data e período por meio de regras automáticas.",
        "Registrar respostas e comentários administrativos para rastreabilidade.",
        "Manter professores informados por avisos, mensagens e notificações de reserva respondida.",
    ])

    for heading, table_data, widths in [
        ("3. Requisitos funcionais", FUNCTIONAL_REQUIREMENTS, [0.55 * inch, 2.1 * inch, 3.85 * inch]),
        ("4. Requisitos não funcionais", NON_FUNCTIONAL_REQUIREMENTS, [0.65 * inch, 1.9 * inch, 3.95 * inch]),
        ("5. Casos de uso", USE_CASES, [0.55 * inch, 1.65 * inch, 1.25 * inch, 4.05 * inch]),
    ]:
        story.append(p(heading, styles["H1Custom"]))
        add_pdf_table(story, styles, table_data, widths)

    story.append(p("6. Arquitetura do sistema", styles["H1Custom"]))
    add_pdf_bullets(story, styles, ARCHITECTURE_ITEMS)
    story.append(p("6.1 Mapa de rotas", styles["H2Custom"]))
    add_pdf_table(story, styles, ROUTES, [1.25 * inch, 2.55 * inch, 2.7 * inch])

    story.append(p("7. Modelo de banco de dados", styles["H1Custom"]))
    story.append(p("O banco está configurado para MySQL/MariaDB, com conexão mysql, host 127.0.0.1, porta 3306 e database reserva_salas. Credenciais sensíveis não foram reproduzidas neste documento.", styles["BodyCustom"]))
    add_pdf_table(story, styles, DATABASE_MODEL, [1.25 * inch, 2.95 * inch, 2.3 * inch])
    story.append(p("Observação técnica: existem migrations duplicadas ou evolutivas para mensagens diretas e cor de blocos. Algumas migrations usam verificações Schema::hasTable e Schema::hasColumn para evitar erro quando a estrutura já existe.", styles["BodyCustom"]))

    story.append(p("8. Telas implementadas", styles["H1Custom"]))
    add_pdf_table(story, styles, SCREENS, [0.95 * inch, 1.25 * inch, 2.25 * inch, 2.05 * inch])

    story.append(p("9. Funcionalidades implementadas", styles["H1Custom"]))
    add_pdf_bullets(story, styles, IMPLEMENTED_FEATURES)

    story.append(p("10. Tecnologias utilizadas", styles["H1Custom"]))
    add_pdf_table(story, styles, TECHNOLOGIES, [1.15 * inch, 1.75 * inch, 3.6 * inch])

    story.append(p("11. Fluxos principais", styles["H1Custom"]))
    for sub, items in [
        ("11.1 Fluxo de reserva simples", [
            "Professor acessa Solicitar Reserva.",
            "Seleciona bloco, sala, data e período.",
            "Sistema consulta disponibilidade por sala, data e período.",
            "Se houver manutenção, a reserva é cancelada com motivo administrativo.",
            "Se houver conflito em solicitação simples, a reserva é cancelada por ocupação.",
            "Se a data for o mesmo dia, a reserva é aprovada automaticamente; caso contrário, fica pendente.",
        ]),
        ("11.2 Fluxo de reserva recorrente", [
            "Professor ativa o modo recorrente e escolhe datas base.",
            "Sistema gera datas semanais até três meses a partir de cada data base.",
            "Todas as reservas recebem o mesmo grupo_recorrencia.",
            "Datas com conflito entram em análise, e datas sem conflito seguem pendentes ou aprovadas conforme regra.",
            "Administrador pode analisar cada data individualmente ou aprovar/recusar todas no fim da lista do grupo.",
        ]),
        ("11.3 Fluxo de manutenção", [
            "Administrador ativa manutenção em bloco ou sala.",
            "Define prazo final ou tempo indeterminado e aviso opcional.",
            "Sistema cancela reservas futuras pendentes, em análise ou aprovadas dentro do período afetado.",
            "Novas solicitações para salas em manutenção recebem mensagem de indisponibilidade.",
        ]),
    ]:
        story.append(p(sub, styles["H2Custom"]))
        add_pdf_bullets(story, styles, items)

    story.append(p("12. Considerações técnicas e oportunidades de melhoria", styles["H1Custom"]))
    add_pdf_bullets(story, styles, [
        "Separar responsabilidades do AdminController em controllers específicos para usuários, infraestrutura, reservas, avisos e mensagens.",
        "Criar Form Requests para centralizar validações e mensagens de erro.",
        "Adicionar foreign keys explícitas para remetente_id e destinatario_id em mensagens_diretas.",
        "Normalizar status de reserva para enum ou constantes de domínio, evitando variações como recusada e rejeitada.",
        "Expandir testes automatizados para autenticação, reserva recorrente, manutenção, conflitos e notificações.",
        "Revisar migrations duplicadas para reduzir risco de inconsistência em ambientes novos.",
        "Adicionar migrations ou jobs de limpeza dedicados em vez de acionar limpeza apenas ao acessar telas.",
    ])

    story.append(p("13. Conclusão", styles["H1Custom"]))
    story.append(p("O sistema reserva-salas apresenta uma base funcional completa para gestão institucional de reservas, contemplando autenticação, perfis, infraestrutura, solicitações simples e recorrentes, aprovação administrativa, manutenção, avisos, mensagens e histórico. A aplicação utiliza padrões reconhecidos do Laravel e organiza a interface em telas Blade especializadas por área.", styles["BodyCustom"]))
    story.append(p("Do ponto de vista técnico, o projeto já cobre as principais necessidades operacionais do domínio. As melhorias recomendadas concentram-se em modularização do controller principal, fortalecimento das constraints de banco, padronização de status e ampliação da cobertura de testes, medidas que aumentariam a manutenção e a confiabilidade em produção.", styles["BodyCustom"]))

    doc.build(story)


if __name__ == "__main__":
    OUT_DIR.mkdir(exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
